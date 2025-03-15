import docx
import shutil
import datetime

def process_document(input_path, output_path, replacement_text, client_name, key_account_name):
    """
    Copies a Word document, replaces the first "Název Projektu" heading, and saves the modified document.

    Args:
        input_path (str): Path to the input Word document.
        output_path (str): Path to save the modified document.
        replacement_text (str): The text to replace "Název Projektu" with.
        client_name (str): The name of the client.
        key_account_name (str): The name of the key account.
    """
    try:
        # Copy the document
        shutil.copy(input_path, output_path)

        # Open the copied document
        doc = docx.Document(output_path)

       # Find and replace the heading
        for paragraph in doc.paragraphs:
            if "NÁZEV PROJEKTU" in paragraph.text:
                original_style = paragraph.style
                paragraph.text = replacement_text
                paragraph.style = original_style  # Keep original style
                break  # Replace only the first occurrence

        # Find and replace the program name
        for paragraph in doc.paragraphs:
            if "Program, ve kterém je (PINKway, ...) - in" in paragraph.text:
                paragraph.text = f"Program: {program_name}"
                paragraph.style = paragraph.style  # Keep original style
                break

        # Find and replace the client name
        for paragraph in doc.paragraphs:
            if "Klient:" in paragraph.text:
                paragraph.text = f"Klient: {client_name}"
                paragraph.style = paragraph.style  # Keep original style
                break

        # Find and replace the key account name
        for paragraph in doc.paragraphs:
            if "Key Account:" in paragraph.text:
                paragraph.text = f"Key Account: {key_account_name}"
                paragraph.style = paragraph.style  # Keep original style
                break

        # Find and replace the date
        current_date = datetime.datetime.now().strftime("%d/%m/%Y")
        for paragraph in doc.paragraphs:
            if "Datum: " in paragraph.text:
                paragraph.text = f"Datum: {current_date}"
                paragraph.style = paragraph.style  # Keep original style
                break

        # Save the modified document
        doc.save(output_path)

        return True, None  # Indicate success

    except Exception as e:
        return False, str(e)  # Indicate failure and return the error message


if __name__ == "__main__":
    input_file = "vzor.docx.docx"

    # Get user input for project name
    replacement = input("Zadejte nový název projektu: ")

    # Get user input for program name
    program_name = input("Zadejte název programu: ")

    # Get user input for client name
    client_name = input("Zadejte jméno klienta: ")

    # Get user input for key account name
    key_account_name = input("Zadejte Vaše jméno: ")

    # Create the output file name
    output_file = f"{replacement}_{program_name}.docx"

    success, error_message = process_document(input_file, output_file, replacement, client_name, key_account_name)

    if success:
        print(f"Dokument '{output_file}' byl úspěšně vytvořen.")
    else:
        print(f"Došlo k chybě: {error_message}")
